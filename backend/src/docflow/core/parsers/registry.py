from __future__ import annotations

import mimetypes
import re
from email import policy
from email.header import decode_header, make_header
from email.message import Message
from email.parser import BytesParser
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from docflow.schemas import DocumentSegment, DocumentType, NormalizedDocument, SegmentKind
from docflow.utils.hashing import sha256_text, short_hash, simhash64
from docflow.utils.text import detect_language, normalize_text
from docflow.utils.time import utc_now


class ParserRegistry:
    def parse_file(
        self, path: Path, tenant_id: str, source: str | None = None
    ) -> NormalizedDocument:
        extension = path.suffix.lower()
        if extension == ".pdf":
            return self._parse_pdf(path, tenant_id, source)
        if extension in {".docx", ".doc"}:
            return self._parse_docx(path, tenant_id, source)
        if extension in {".md", ".markdown"}:
            return self._parse_markdown(path, tenant_id, source)
        if extension in {".eml", ".msg"}:
            return self._parse_email(path, tenant_id, source)
        if extension in {".html", ".htm"}:
            return self._parse_html(path, tenant_id, source)
        if extension in {".log"}:
            return self._parse_log(path, tenant_id, source)
        return self._parse_text(path, tenant_id, source)

    def _parse_pdf(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        reader = PdfReader(str(path))
        segments: list[DocumentSegment] = []
        pages: list[str] = []
        position = 0
        for page_number, page in enumerate(reader.pages, start=1):
            text = normalize_text(page.extract_text() or "")
            pages.append(text)
            for paragraph in [item for item in text.split("\n\n") if item.strip()]:
                position += 1
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.paragraph,
                        text=paragraph.strip(),
                        position=position,
                        page=page_number,
                    )
                )
        title = reader.metadata.title if reader.metadata else None
        metadata = {
            "page_count": len(reader.pages),
            "source_mime": mimetypes.guess_type(path.name)[0] or "application/pdf",
            "pdf_metadata": {
                "author": getattr(reader.metadata, "author", None) if reader.metadata else None,
                "producer": getattr(reader.metadata, "producer", None) if reader.metadata else None,
            },
        }
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.pdf,
            source=source,
            title=title,
            raw_text="\n\n".join(pages),
            segments=segments,
            metadata=metadata,
        )

    def _parse_docx(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        document = DocxDocument(str(path))
        segments: list[DocumentSegment] = []
        raw_parts: list[str] = []
        position = 0
        title = None
        for paragraph in document.paragraphs:
            text = normalize_text(paragraph.text)
            if not text:
                continue
            raw_parts.append(text)
            position += 1
            kind = (
                SegmentKind.heading
                if paragraph.style.name.lower().startswith("heading")
                else SegmentKind.paragraph
            )
            level = (
                self._heading_level(paragraph.style.name) if kind == SegmentKind.heading else None
            )
            if not title and kind == SegmentKind.heading:
                title = text
            segments.append(
                DocumentSegment(
                    segment_id=f"seg-{position}",
                    kind=kind,
                    text=text,
                    level=level,
                    position=position,
                )
            )
        for table in document.tables:
            for row in table.rows:
                text = " | ".join(
                    normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)
                )
                if not text:
                    continue
                raw_parts.append(text)
                position += 1
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.table_row,
                        text=text,
                        position=position,
                    )
                )
        metadata = {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "core_properties": {
                "author": document.core_properties.author,
                "created": document.core_properties.created.isoformat()
                if document.core_properties.created
                else None,
                "modified": document.core_properties.modified.isoformat()
                if document.core_properties.modified
                else None,
            },
        }
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.docx,
            source=source,
            title=title,
            raw_text="\n\n".join(raw_parts),
            segments=segments,
            metadata=metadata,
        )

    def _parse_markdown(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        normalized = normalize_text(text)
        segments: list[DocumentSegment] = []
        position = 0
        title = None
        for line in normalized.splitlines():
            if not line.strip():
                continue
            position += 1
            if line.startswith("#"):
                heading_level = len(line) - len(line.lstrip("#"))
                content = line.lstrip("#").strip()
                if not title:
                    title = content
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.heading,
                        text=content,
                        level=heading_level,
                        position=position,
                    )
                )
            elif line.startswith(("- ", "* ")):
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.list_item,
                        text=line[2:].strip(),
                        position=position,
                    )
                )
            elif line.startswith(">"):
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.quote,
                        text=line[1:].strip(),
                        position=position,
                    )
                )
            else:
                segments.append(
                    DocumentSegment(
                        segment_id=f"seg-{position}",
                        kind=SegmentKind.paragraph,
                        text=line.strip(),
                        position=position,
                    )
                )
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.markdown,
            source=source,
            title=title,
            raw_text=normalized,
            segments=segments,
            metadata={"source_mime": "text/markdown"},
        )

    def _parse_email(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        with path.open("rb") as handle:
            message = BytesParser(policy=policy.default).parse(handle)
        subject = self._decode_header_value(message.get("Subject"))
        sender = self._decode_header_value(message.get("From"))
        recipients = self._decode_header_value(message.get("To"))
        sent_at = self._decode_header_value(message.get("Date"))
        body = self._extract_email_body(message)
        normalized_body = normalize_text(body)
        segments = [
            DocumentSegment(
                segment_id="seg-1", kind=SegmentKind.metadata, text=f"From: {sender}", position=1
            ),
            DocumentSegment(
                segment_id="seg-2", kind=SegmentKind.metadata, text=f"To: {recipients}", position=2
            ),
            DocumentSegment(
                segment_id="seg-3", kind=SegmentKind.metadata, text=f"Date: {sent_at}", position=3
            ),
            DocumentSegment(
                segment_id="seg-4",
                kind=SegmentKind.heading,
                text=subject or "Untitled email",
                level=1,
                position=4,
            ),
        ]
        position = 4
        for paragraph in [item for item in normalized_body.split("\n\n") if item.strip()]:
            position += 1
            segments.append(
                DocumentSegment(
                    segment_id=f"seg-{position}",
                    kind=SegmentKind.paragraph,
                    text=paragraph,
                    position=position,
                )
            )
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.email,
            source=source,
            title=subject,
            raw_text=normalized_body,
            segments=segments,
            metadata={
                "email_headers": {"from": sender, "to": recipients, "date": sent_at},
                "thread_hint": self._guess_email_thread_id(message),
            },
        )

    def _parse_html(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        title = normalize_text(soup.title.text) if soup.title else None
        segments: list[DocumentSegment] = []
        raw_parts: list[str] = []
        position = 0
        for element in soup.find_all(["h1", "h2", "h3", "p", "li", "blockquote", "tr"]):
            text = normalize_text(element.get_text(" ", strip=True))
            if not text:
                continue
            raw_parts.append(text)
            position += 1
            if element.name in {"h1", "h2", "h3"}:
                kind = SegmentKind.heading
                level = int(element.name[1])
            elif element.name == "li":
                kind = SegmentKind.list_item
                level = None
            elif element.name == "blockquote":
                kind = SegmentKind.quote
                level = None
            elif element.name == "tr":
                kind = SegmentKind.table_row
                level = None
            else:
                kind = SegmentKind.paragraph
                level = None
            segments.append(
                DocumentSegment(
                    segment_id=f"seg-{position}",
                    kind=kind,
                    text=text,
                    level=level,
                    position=position,
                )
            )
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.html,
            source=source,
            title=title,
            raw_text="\n\n".join(raw_parts),
            segments=segments,
            metadata={"source_mime": "text/html"},
        )

    def _parse_log(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        normalized = normalize_text(text)
        segments: list[DocumentSegment] = []
        position = 0
        for line in normalized.splitlines():
            if not line.strip():
                continue
            position += 1
            segments.append(
                DocumentSegment(
                    segment_id=f"seg-{position}",
                    kind=SegmentKind.line,
                    text=line.strip(),
                    position=position,
                )
            )
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=DocumentType.log,
            source=source,
            title=path.stem,
            raw_text=normalized,
            segments=segments,
            metadata={"source_mime": "text/plain", "log_like": True},
        )

    def _parse_text(self, path: Path, tenant_id: str, source: str | None) -> NormalizedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        normalized = normalize_text(text)
        inferred_type = self._infer_text_type(normalized, path)
        segments: list[DocumentSegment] = []
        position = 0
        title = None
        paragraphs = (
            normalized.split("\n\n")
            if inferred_type != DocumentType.log
            else normalized.splitlines()
        )
        for block in paragraphs:
            block = block.strip()
            if not block:
                continue
            position += 1
            kind = SegmentKind.paragraph
            if inferred_type == DocumentType.chat and re.match(
                r"^\[?\d{4}[-/]\d{2}[-/]\d{2}", block
            ):
                kind = SegmentKind.line
            if position == 1 and len(block) < 120:
                title = block
            segments.append(
                DocumentSegment(
                    segment_id=f"seg-{position}",
                    kind=kind,
                    text=block,
                    position=position,
                )
            )
        return self._build_document(
            path=path,
            tenant_id=tenant_id,
            doc_type=inferred_type,
            source=source,
            title=title or path.stem,
            raw_text=normalized,
            segments=segments,
            metadata={"source_mime": mimetypes.guess_type(path.name)[0] or "text/plain"},
        )

    def _build_document(
        self,
        *,
        path: Path,
        tenant_id: str,
        doc_type: DocumentType,
        source: str | None,
        title: str | None,
        raw_text: str,
        segments: list[DocumentSegment],
        metadata: dict,
    ) -> NormalizedDocument:
        now = utc_now()
        normalized_text = normalize_text(raw_text)
        checksum = sha256_text(raw_text)
        normalized_checksum = sha256_text(normalized_text)
        document_id = short_hash(f"{tenant_id}:{path.resolve()}:{normalized_checksum}", length=16)
        return NormalizedDocument(
            document_id=document_id,
            source=source or str(path.resolve()),
            filename=path.name,
            doc_type=doc_type,
            tenant_id=tenant_id,
            title=title,
            language=detect_language(normalized_text),
            checksum=checksum,
            normalized_checksum=normalized_checksum,
            simhash=simhash64(normalized_text),
            raw_text=raw_text,
            normalized_text=normalized_text,
            segments=segments,
            metadata=metadata,
            created_at=now,
            updated_at=now,
        )

    @staticmethod
    def _decode_header_value(value: str | None) -> str:
        if value is None:
            return ""
        decoded = str(make_header(decode_header(str(value))))
        return normalize_text(decoded)

    @staticmethod
    def _extract_email_body(message: Message) -> str:
        if message.is_multipart():
            preferred: str | None = None
            for part in message.walk():
                content_type = part.get_content_type()
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
                charset = part.get_content_charset() or "utf-8"
                decoded = payload.decode(charset, errors="ignore")
                if content_type == "text/plain":
                    return decoded
                if content_type == "text/html" and preferred is None:
                    preferred = BeautifulSoup(decoded, "html.parser").get_text("\n")
            return preferred or ""
        payload = message.get_payload(decode=True)
        if payload:
            return payload.decode(message.get_content_charset() or "utf-8", errors="ignore")
        return ""

    @staticmethod
    def _guess_email_thread_id(message: Message) -> str | None:
        return (
            message.get("Thread-Index") or message.get("Message-ID") or message.get("In-Reply-To")
        )

    @staticmethod
    def _heading_level(style_name: str) -> int:
        match = re.search(r"(\d+)$", style_name)
        return int(match.group(1)) if match else 1

    @staticmethod
    def _infer_text_type(text: str, path: Path) -> DocumentType:
        if re.search(r"^\[\d{4}-\d{2}-\d{2}.*?\]\s*[^:]+:", text, re.MULTILINE):
            return DocumentType.chat
        if re.search(r"^\d{4}-\d{2}-\d{2}.*?\b(INFO|WARN|ERROR|DEBUG)\b", text, re.MULTILINE):
            return DocumentType.log
        if path.suffix.lower() in {".txt"}:
            return DocumentType.txt
        return DocumentType.unknown
